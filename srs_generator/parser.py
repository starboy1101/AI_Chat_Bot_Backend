from __future__ import annotations

import html
import re
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from srs_generator.models import DocumentBlock, ParsedDocument
from srs_generator.utils import normalize_space


class DocumentParser:
    """Parse engineering input documents into ordered text/table blocks."""

    def parse_path(self, path: str | Path) -> ParsedDocument:
        p = Path(path)
        return self.parse_bytes(p.read_bytes(), p.name)

    def parse_bytes(self, file_bytes: bytes, source_name: str) -> ParsedDocument:
        suffix = Path(source_name).suffix.lower()
        if suffix == ".docx":
            return self._parse_docx(file_bytes, source_name)
        if suffix == ".pdf":
            return self._parse_pdf(file_bytes, source_name)
        if suffix == ".txt":
            return self._parse_txt(file_bytes, source_name)
        if suffix == ".doc":
            return self._parse_legacy_doc(file_bytes, source_name)
        raise ValueError(f"Unsupported document type: {suffix or source_name}")

    def _parse_docx(self, file_bytes: bytes, source_name: str) -> ParsedDocument:
        document = Document(BytesIO(file_bytes))
        blocks: List[DocumentBlock] = []
        text_parts: List[str] = []

        for idx, item in enumerate(self._iter_docx_blocks(document)):
            if isinstance(item, Paragraph):
                text = normalize_space(item.text)
                if not text:
                    continue
                style_name = item.style.name if item.style is not None else None
                level = self._heading_level(style_name)
                blocks.append(
                    DocumentBlock(
                        kind="heading" if level else "paragraph",
                        text=text,
                        heading_level=level,
                        style=style_name,
                        index=idx,
                    )
                )
                text_parts.append(text)
            elif isinstance(item, Table):
                rows = [[normalize_space(cell.text) for cell in row.cells] for row in item.rows]
                row_text = "\n".join(" | ".join(cell for cell in row if cell) for row in rows)
                blocks.append(
                    DocumentBlock(
                        kind="table",
                        text=normalize_space(row_text),
                        rows=rows,
                        index=idx,
                    )
                )
                if row_text.strip():
                    text_parts.append(row_text)

        return ParsedDocument(
            source_name=source_name,
            source_type="docx",
            text=normalize_space("\n\n".join(text_parts)),
            blocks=blocks,
            metadata={"tables": len(document.tables), "paragraphs": len(document.paragraphs)},
        )

    def _parse_pdf(self, file_bytes: bytes, source_name: str) -> ParsedDocument:
        page_texts: List[str] = []
        try:
            import fitz

            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page_index, page in enumerate(doc):
                    page_texts.append(page.get_text("text") or "")
        except Exception:
            try:
                import pdfplumber

                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        page_texts.append(page.extract_text() or "")
            except Exception as exc:
                raise ValueError(f"PDF unreadable: {exc}") from exc

        blocks = [
            DocumentBlock(kind="paragraph", text=normalize_space(text), page=i + 1, index=i)
            for i, text in enumerate(page_texts)
            if normalize_space(text)
        ]
        return ParsedDocument(
            source_name=source_name,
            source_type="pdf",
            text=normalize_space("\n\n".join(page_texts)),
            blocks=blocks,
            metadata={"pages": len(page_texts)},
        )

    def _parse_txt(self, file_bytes: bytes, source_name: str) -> ParsedDocument:
        decoded = self._decode_text(file_bytes)
        blocks = []
        for idx, para in enumerate(re.split(r"\n\s*\n", decoded)):
            text = normalize_space(para)
            if text:
                blocks.append(DocumentBlock(kind="paragraph", text=text, index=idx))
        return ParsedDocument(source_name=source_name, source_type="txt", text=normalize_space(decoded), blocks=blocks)

    def _parse_legacy_doc(self, file_bytes: bytes, source_name: str) -> ParsedDocument:
        chunks = re.findall(rb"[A-Za-z0-9][\x20-\x7E]{5,}", file_bytes)
        if not chunks:
            raise ValueError("DOC unreadable or empty")
        text = normalize_space(" ".join(chunk.decode("latin-1", errors="ignore") for chunk in chunks))
        blocks = [DocumentBlock(kind="paragraph", text=text, index=0)]
        return ParsedDocument(source_name=source_name, source_type="doc", text=text, blocks=blocks)

    def _decode_text(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8", "utf-16", "utf-16le", "utf-16be", "latin-1"):
            try:
                text = file_bytes.decode(encoding)
                if normalize_space(text):
                    return text
            except Exception:
                continue
        return file_bytes.decode("utf-8", errors="ignore")

    def _iter_docx_blocks(self, document: DocxDocument) -> Iterable[Paragraph | Table]:
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, document)
            elif child.tag.endswith("}tbl"):
                yield Table(child, document)

    def _heading_level(self, style_name: str | None) -> int | None:
        if not style_name:
            return None
        match = re.search(r"heading\s+(\d+)", style_name, flags=re.I)
        if match:
            return int(match.group(1))
        return None


def extract_docx_text_fast(file_bytes: bytes) -> str:
    """Fallback text extractor that reads raw DOCX XML parts."""
    with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
        names = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
        parts: List[str] = []
        for name in names:
            xml = archive.read(name).decode("utf-8", errors="ignore")
            xml = re.sub(r"</w:p>", "\n", xml)
            xml = re.sub(r"</w:tr>", "\n", xml)
            xml = re.sub(r"<[^>]+>", " ", xml)
            plain = normalize_space(html.unescape(xml))
            if plain:
                parts.append(plain)
    return normalize_space("\n\n".join(parts))

